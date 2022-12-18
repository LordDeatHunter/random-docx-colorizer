import random

from docx import Document
from docx.shared import RGBColor

document = Document(input("Filename (with extension: "))
print('Doing magic...')
color_options = [
    (0xAE, 0x00, 0x00),
    (0x00, 0xAE, 0x00),
    (0x00, 0x00, 0xAE)
]

paragraphs = []
color_count = 0
for p in document.paragraphs:
    runs = []
    styles = []
    for r in p.runs:
        text = r.text.split(' ')
        color_count += len(text)
        runs.append(text)
        styles.append(r.font)
    paragraphs.append((runs, styles))

colors = []
for color in range(color_count):
    while True:
        color = random.choice(color_options)
        if len(color_options) < 2 or len(colors) == 0 or color != colors[-1]:
            colors.append(color)
            break

for i, p in enumerate(paragraphs):
    document.paragraphs[i].clear()
    for runs, s in zip(p[0], p[1]):
        for run in runs:
            document.paragraphs[i].add_run(f'{run} ')
            document.paragraphs[i].runs[-1].font.color.rgb = RGBColor(
                *colors.pop(0)
            )
            # TODO: figure out how to copy the style directly
            document.paragraphs[i].runs[-1].font.all_caps = s.all_caps
            document.paragraphs[i].runs[-1].font.bold = s.bold
            document.paragraphs[i].runs[-1].font.italic = s.italic
            document.paragraphs[i].runs[-1].font.cs_bold = s.cs_bold
            document.paragraphs[i].runs[-1].font.cs_italic = s.cs_italic
            document.paragraphs[i].runs[-1].font.double_strike = s.double_strike
            document.paragraphs[i].runs[-1].font.emboss = s.emboss
            document.paragraphs[i].runs[-1].font.hidden = s.hidden
            document.paragraphs[i].runs[-1].font.name = s.name
            document.paragraphs[i].runs[-1].font.highlight_color = s.highlight_color
            document.paragraphs[i].runs[-1].font.outline = s.outline
            document.paragraphs[i].runs[-1].font.shadow = s.shadow
            document.paragraphs[i].runs[-1].font.rtl = s.rtl
            document.paragraphs[i].runs[-1].font.size = s.size
            document.paragraphs[i].runs[-1].font.no_proof = s.no_proof
            document.paragraphs[i].runs[-1].font.math = s.math
            document.paragraphs[i].runs[-1].font.imprint = s.imprint
            document.paragraphs[i].runs[-1].font.complex_script = s.complex_script
            document.paragraphs[i].runs[-1].font.small_caps = s.small_caps
            document.paragraphs[i].runs[-1].font.snap_to_grid = s.snap_to_grid
            document.paragraphs[i].runs[-1].font.spec_vanish = s.spec_vanish
            document.paragraphs[i].runs[-1].font.strike = s.strike
            document.paragraphs[i].runs[-1].font.subscript = s.subscript
            document.paragraphs[i].runs[-1].font.superscript = s.superscript
            document.paragraphs[i].runs[-1].font.underline = s.underline
            document.paragraphs[i].runs[-1].font.web_hidden = s.web_hidden

document.save(f'{input("Output filename (without extension): ")}.docx')
